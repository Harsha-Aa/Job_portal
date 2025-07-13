import streamlit as st
import PyPDF2
import docx
import re
import io
from typing import Optional

class ResumeParser:
    """
    Resume parser that extracts text from PDF and DOCX files
    """
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx']
    
    def parse_resume(self, uploaded_file) -> Optional[str]:
        """
        Parse uploaded resume file and extract text
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            str: Extracted text from resume or None if parsing fails
        """
        try:
            file_extension = uploaded_file.name.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                return self._parse_pdf(uploaded_file)
            elif file_extension == 'docx':
                return self._parse_docx(uploaded_file)
            else:
                st.error(f"Unsupported file format: {file_extension}")
                return None
                
        except Exception as e:
            st.error(f"Error parsing resume: {str(e)}")
            return None
    
    def _parse_pdf(self, uploaded_file) -> Optional[str]:
        """
        Extract text from PDF file
        
        Args:
            uploaded_file: PDF file object
            
        Returns:
            str: Extracted text or None if parsing fails
        """
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            
            # Create PDF reader
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            
            # Extract text from all pages
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            # Clean up text
            text = self._clean_text(text)
            
            return text if text.strip() else None
            
        except Exception as e:
            st.error(f"Error parsing PDF: {str(e)}")
            return None
    
    def _parse_docx(self, uploaded_file) -> Optional[str]:
        """
        Extract text from DOCX file
        
        Args:
            uploaded_file: DOCX file object
            
        Returns:
            str: Extracted text or None if parsing fails
        """
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            
            # Read DOCX file
            doc = docx.Document(uploaded_file)
            
            # Extract text from all paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # Clean up text
            text = self._clean_text(text)
            
            return text if text.strip() else None
            
        except Exception as e:
            st.error(f"Error parsing DOCX: {str(e)}")
            return None
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        if not text:
            return ""
        
        # Remove extra whitespace and normalize line breaks
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n', text)
        
        # Remove special characters but keep alphanumeric and common punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)@]', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    def extract_contact_info(self, text: str) -> dict:
        """
        Extract contact information from resume text
        
        Args:
            text: Resume text
            
        Returns:
            dict: Extracted contact information
        """
        contact_info = {}
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact_info['email'] = emails[0]
        
        # Extract phone numbers
        phone_pattern = r'(\+?1?[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'
        phones = re.findall(phone_pattern, text)
        if phones:
            contact_info['phone'] = ''.join(phones[0])
        
        # Extract LinkedIn profile
        linkedin_pattern = r'linkedin\.com/in/[A-Za-z0-9\-]+'
        linkedin = re.findall(linkedin_pattern, text)
        if linkedin:
            contact_info['linkedin'] = linkedin[0]
        
        return contact_info
    
    def extract_sections(self, text: str) -> dict:
        """
        Extract different sections from resume text
        
        Args:
            text: Resume text
            
        Returns:
            dict: Extracted sections
        """
        sections = {}
        
        # Common section headers
        section_patterns = {
            'education': r'(?i)(education|academic|university|college|degree)',
            'experience': r'(?i)(experience|employment|work|career|professional)',
            'skills': r'(?i)(skills|technical|proficiencies|competencies)',
            'projects': r'(?i)(projects|portfolio|work samples)',
            'certifications': r'(?i)(certifications|certificates|credentials)',
            'achievements': r'(?i)(achievements|awards|accomplishments|honors)'
        }
        
        # Split text into lines
        lines = text.split('\n')
        current_section = None
        section_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line is a section header
            found_section = None
            for section_name, pattern in section_patterns.items():
                if re.search(pattern, line):
                    found_section = section_name
                    break
            
            if found_section:
                # Save previous section
                if current_section and section_content:
                    sections[current_section] = '\n'.join(section_content)
                
                # Start new section
                current_section = found_section
                section_content = []
            elif current_section:
                section_content.append(line)
        
        # Save last section
        if current_section and section_content:
            sections[current_section] = '\n'.join(section_content)
        
        return sections
    
    def get_resume_summary(self, text: str) -> dict:
        """
        Generate a summary of the resume
        
        Args:
            text: Resume text
            
        Returns:
            dict: Resume summary
        """
        if not text:
            return {}
        
        summary = {
            'word_count': len(text.split()),
            'character_count': len(text),
            'estimated_pages': max(1, len(text) // 3000),  # Rough estimate
            'sections': self.extract_sections(text),
            'contact_info': self.extract_contact_info(text)
        }
        
        return summary
